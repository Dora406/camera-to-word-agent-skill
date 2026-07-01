import cv2
import requests
import json
import os
import base64
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 获取当前电脑的桌面路径
desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")

def create_professional_docx(text_content):
    """根据AI生成的结构化内容，在桌面创建一份精排版的Word文档"""
    if not text_content.strip():
        print(" [错误] AI返回内容为空，无法生成Word文档！")
        return False
        
    doc = Document()
    
    # 设置页面边距（1英寸）
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = text_content.split('\n')
    has_content = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        has_content = True
        p = doc.add_paragraph()
        
        # 智能排版逻辑
        if line.startswith('# ') or line.startswith('TITLE:') or line.startswith('Title:'):
            clean_line = line.replace('#', '').replace('TITLE:', '').replace('Title:', '').strip()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_line)
            run.font.name = '微软雅黑'
            run.font.size = Pt(18)
            run.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = '宋体'
            run.font.size = Pt(12)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.space_after = Pt(6)

    if has_content:
        timestamp = time.strftime("%H%M%S")
        safe_output_path = os.path.join(desktop_path, f"Dell_AI_Transformed_{timestamp}.docx")
        
        try:
            doc.save(safe_output_path)
            print(f" [排版成功] Word文档已生成：\n {safe_output_path}")
            if os.name == 'nt':
                os.system(f'start "" "{safe_output_path}"')
            return True
        except Exception as file_err:
            print(f" [文件保存失败] 错误详情: {file_err}")
            return False
    else:
        print(" [警告] 未提取到有效文字，未生成Word。")
        return False

def analyze_image(image_path):
    """将图片发送给本地Ollama运行的 Qwen3-VL 4B 进行轻量高质分析"""
    print(f"\n [OpenClaw 激活] 正在调用 Qwen3-VL (4B) 极速模型解析图像...")
    print("💡 提示：4B模型具备极高的响应速度，戴尔工作站正在进行秒级本地推理！")
    
    with open(image_path, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
    
    # 核心修改：已替换为 4B 模型
    MODEL_NAME = "qwen3-vl:4b" 
    prompt_text = (
        "你是一个极其专业的智能文档扫描特工。请精确提取这张图片中的所有文字内容。"
        "必须保持原有的段落结构。如果是标题，请在行首加上 '# '。"
        "直接输出提取后的文字，绝对不要包含任何多余的解释、寒暄。"
    )
    
    try:
        # 4B 模型速度极快，超时时间保留 60 秒绰绰有余
        response = requests.post('http://localhost:11434/api/generate', json={
            "model": MODEL_NAME, 
            "prompt": prompt_text,
            "stream": False,
            "images": [encoded_string]
        }, timeout=60)
        
        ai_response = response.json().get('response', '').strip()
        
        print("\n---  [Qwen3-VL:4b 实时输出] ---")
        print(ai_response)
        print("---------------------------------\n")
        
        # 执行排版并返回结果状态
        return create_professional_docx(ai_response)
        
    except Exception as ollama_err:
        print(f"[API错误] 无法连接到Ollama: {ollama_err}")
        return False

# ---- 主程序开始 ----
cap = cv2.VideoCapture(0)
print("=== 戴尔Precision × OpenClaw 极速智能工作站已就绪 ===")
print("💡 快捷键提示: [Space] 拍照扫描 | [L] 读取桌面 test.jpg | [Esc] 退出程序")

while True:
    ret, frame = cap.read()
    if not ret:
        break
        
    clean_frame = frame.copy()
    h, w, _ = frame.shape
    cv2.rectangle(frame, (int(w*0.15), int(h*0.1)), (int(w*0.85), int(h*0.9)), (0, 255, 0), 2)
    cv2.putText(frame, "DELL PRECISION - QWEN3-VL:4B ACTIVE", (20, 30), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0, 255, 0), 2)
    cv2.putText(frame, "SPACE: Scan & Exit | L: Test Img | Esc: Exit", (20, h - 20), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (255, 255, 255), 1)
    
    cv2.imshow('OpenClaw Camera Agent (4B Speed Edition)', frame)
    key = cv2.waitKey(1) & 0xFF
    
    # 情况A：用户按下空格，拍照
    if key == ord(' '):
        img_path = 'live_clean_snapshot.jpg'
        cv2.imwrite(img_path, clean_frame)
        success = analyze_image(img_path)
        break
        
    # 情况B：按下L键，调用桌面的备份高清图
    elif key == ord('l') or key == ord('L'):
        mock_img_path = os.path.join(desktop_path, "test.jpg")
        if os.path.exists(mock_img_path):
            analyze_image(mock_img_path)
        else:
            print(f" 未在桌面找到 'test.jpg' 文件。")
        break
            
    # 情况C：按下 Esc 键手动退出
    elif key == 27:
        print(" [提示] 用户按下 Esc 键，正在手动退出并释放资源...")
        break

# 彻底销毁资源，确保后台干净
cap.release()
cv2.destroyAllWindows()
print("===  任务完成，戴尔本地硬件资源已完全释放 ===")